# =================================================================================
#  ФАЙЛ: bot.py (V5.9 - С ГИБКИМ УПРАВЛЕНИЕМ ДОСТУПОМ)
# =================================================================================

# --- 1. ИМПОРТЫ ---
import os
import logging
import asyncio
from datetime import datetime
import zipfile
import io
from typing import List, Dict, Any, Optional, Tuple, Set
import psycopg2
import yt_dlp
import telegram
import uuid
import time
import docx
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from telegram import Update, ReplyKeyboardMarkup, Message, InlineKeyboardButton, InlineKeyboardMarkup, ReplyKeyboardRemove, MessageOriginUser
from telegram.ext import (
    Application,
    CommandHandler,
    MessageHandler,
    filters,
    ContextTypes,
    ConversationHandler,
    CallbackQueryHandler,
)
from cryptography import x509
from cryptography.hazmat.backends import default_backend
from openpyxl import Workbook
from openpyxl.styles import PatternFill
from openpyxl.utils import get_column_letter
from dotenv import load_dotenv

load_dotenv()


# --- 2. НАСТРОЙКА И КОНСТАНТЫ ---
TELEGRAM_BOT_TOKEN = os.environ.get("TELEGRAM_BOT_TOKEN")
DATABASE_URL = os.environ.get("DATABASE_URL")
ADMIN_USER_ID = 96238783  # ID главного администратора, который не может быть удален

logging.basicConfig(format='%(asctime)s - %(name)s - %(levelname)s - %(message)s', level=logging.INFO)
logger = logging.getLogger(__name__)

# --- СИСТЕМА РАЗРЕШЕНИЙ (ROLES/PERMISSIONS) ---
AVAILABLE_PERMISSIONS = {
    "cert_analysis": "📜 Анализ сертификатов",
    "akc_form": "📄 Заявка АЦК",
    "youtube": "🎬 Скачивание с YouTube",
    "admin": "👑 Администрирование"
}

def has_permission(user_id: int, feature: str, context: ContextTypes.DEFAULT_TYPE) -> bool:
    """Проверяет, есть ли у пользователя доступ к функции, используя данные из context."""
    permissions_dict = context.bot_data.get('permissions', {})
    user_permissions = permissions_dict.get(user_id, {}).get('perms', set())
    if "admin" in user_permissions:
        return True
    return feature in user_permissions

# -------------------------------------------------

MAX_FILE_SIZE = 20 * 1024 * 1024
MAX_VIDEO_SIZE_BYTES = 49 * 1024 * 1024
EXPIRATION_THRESHOLD_DAYS = 30
RED_FILL = PatternFill(start_color="FFCCCC", end_color="FFCCCC", fill_type="solid")
ORANGE_FILL = PatternFill(start_color="FFDDAA", end_color="FFDDAA", fill_type="solid")
GREEN_FILL = PatternFill(start_color="CCFFCC", end_color="CCFFCC", fill_type="solid")
EXCEL_HEADERS: Tuple[str, ...] = ("ФИО", "Учреждение", "Серийный номер", "Действителен с", "Действителен до", "Осталось дней")
ALLOWED_EXTENSIONS: Tuple[str, ...] = ('.cer', '.crt', '.pem', '.der')
YOUTUBE_URL_PATTERN = r'(https?://)?(www\.)?(youtube|youtu|youtube-nocookie)\.(com|be)/(watch\?v=|embed/|v/|.+\?v=)?([^&=%\?]{11})'

# Состояния для диалогов
(
    AWAITING_YOUTUBE_LINK, CONFIRMING_DOWNLOAD,
    AKC_CONFIRM_DEFAULTS, AKC_SENDER_FIO, AKC_ORG_NAME, AKC_INN_KPP, AKC_MUNICIPALITY,
    AKC_AWAIT_CERTIFICATES, AKC_ROLE, AKC_CITP_NAME, AKC_CONFIRM_LOGINS, AKC_LOGINS, AKC_ACTION,
    CERT_AWAIT_FILES, CERT_AWAIT_THRESHOLD, CERT_TYPING_THRESHOLD,
    ACCESS_MENU, AWAITING_USER_INFO, AWAITING_PERMISSIONS, AWAITING_USER_TO_DELETE
) = range(20)


# --- 3. РАБОТА С БАЗОЙ ДАННЫХ POSTGRESQL ---
def get_db_connection():
    """Устанавливает соединение с базой данных PostgreSQL."""
    try:
        conn = psycopg2.connect(DATABASE_URL)
        return conn
    except Exception as e:
        logger.error(f"Не удалось подключиться к базе данных: {e}")
        return None

def init_database():
    """Инициализирует таблицы в базе данных и гарантирует наличие администратора."""
    conn = get_db_connection()
    if not conn: return
    try:
        with conn.cursor() as cursor:
            cursor.execute('CREATE TABLE IF NOT EXISTS user_settings (user_id BIGINT PRIMARY KEY, threshold INTEGER NOT NULL)')
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS akc_sender_defaults (
                    user_id BIGINT PRIMARY KEY,
                    sender_fio TEXT NOT NULL,
                    org_name TEXT NOT NULL,
                    inn_kpp TEXT NOT NULL,
                    municipality TEXT NOT NULL
                )
            ''')
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS akc_login_defaults (
                    user_id BIGINT PRIMARY KEY,
                    logins TEXT NOT NULL
                )
            ''')
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS user_permissions (
                    user_id BIGINT PRIMARY KEY,
                    username TEXT,
                    permissions TEXT NOT NULL
                )
            ''')
            # Гарантируем, что у главного админа всегда есть права
            cursor.execute(
                """
                INSERT INTO user_permissions (user_id, username, permissions) VALUES (%s, %s, %s) 
                ON CONFLICT (user_id) DO UPDATE SET username = EXCLUDED.username, permissions = EXCLUDED.permissions;
                """,
                (ADMIN_USER_ID, 'Главный Администратор', 'admin')
            )
        conn.commit()
        logger.info("База данных PostgreSQL успешно инициализирована.")
    except Exception as e:
        logger.error(f"Ошибка при инициализации таблиц: {e}")
    finally:
        if conn: conn.close()

def db_load_all_permissions() -> Dict[int, Dict[str, Any]]:
    """Загружает все разрешения из базы данных."""
    conn = get_db_connection()
    if not conn: return {}
    permissions = {}
    try:
        with conn.cursor() as cursor:
            cursor.execute("SELECT user_id, username, permissions FROM user_permissions")
            records = cursor.fetchall()
            for record in records:
                user_id, username, perms_str = record
                permissions[user_id] = {'name': username, 'perms': set(perms_str.split(','))}
    except Exception as e:
        logger.error(f"Ошибка при загрузке разрешений: {e}")
    finally:
        if conn: conn.close()
    return permissions

def db_save_user_permissions(user_id: int, username: str, permissions: Set[str]):
    """Сохраняет или обновляет разрешения для пользователя."""
    conn = get_db_connection()
    if not conn: return
    perms_str = ",".join(sorted(list(permissions)))
    try:
        with conn.cursor() as cursor:
            cursor.execute(
                "INSERT INTO user_permissions (user_id, username, permissions) VALUES (%s, %s, %s) "
                "ON CONFLICT (user_id) DO UPDATE SET username = EXCLUDED.username, permissions = EXCLUDED.permissions;",
                (user_id, username, perms_str)
            )
        conn.commit()
        logger.info(f"Разрешения для пользователя {user_id} ({username}) сохранены.")
    except Exception as e:
        logger.error(f"Ошибка при сохранении разрешений для {user_id}: {e}")
    finally:
        if conn: conn.close()

def db_delete_user(user_id: int):
    """Удаляет пользователя из таблицы разрешений."""
    conn = get_db_connection()
    if not conn: return
    try:
        with conn.cursor() as cursor:
            cursor.execute("DELETE FROM user_permissions WHERE user_id = %s", (user_id,))
        conn.commit()
        logger.info(f"Пользователь {user_id} удален.")
    except Exception as e:
        logger.error(f"Ошибка при удалении пользователя {user_id}: {e}")
    finally:
        if conn: conn.close()

def save_user_threshold(user_id: int, threshold: int):
    """Сохраняет порог дней для пользователя в базу данных."""
    conn = get_db_connection()
    if not conn: return
    try:
        with conn.cursor() as cursor:
            cursor.execute("INSERT INTO user_settings (user_id, threshold) VALUES (%s, %s) ON CONFLICT (user_id) DO UPDATE SET threshold = EXCLUDED.threshold;",(user_id, threshold))
        conn.commit()
    finally:
        if conn: conn.close()

def load_user_threshold(user_id: int) -> Optional[int]:
    """Загружает порог дней для пользователя из базы данных."""
    conn = get_db_connection()
    if not conn: return None
    try:
        with conn.cursor() as cursor:
            cursor.execute("SELECT threshold FROM user_settings WHERE user_id = %s", (user_id,))
            result = cursor.fetchone()
        return result[0] if result else None
    finally:
        if conn: conn.close()

async def get_user_threshold(user_id: int, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Получает порог дней из кэша или базы данных."""
    if 'threshold' in context.user_data:
        return context.user_data['threshold']
    
    threshold_from_db = load_user_threshold(user_id)
    if threshold_from_db is not None:
        context.user_data['threshold'] = threshold_from_db
        return threshold_from_db
    
    return EXPIRATION_THRESHOLD_DAYS

def save_akc_defaults(user_id: int, form_data: dict):
    """Сохраняет данные шапки заявки АЦК как шаблон."""
    conn = get_db_connection()
    if not conn: return
    try:
        with conn.cursor() as cursor:
            cursor.execute(
                "INSERT INTO akc_sender_defaults (user_id, sender_fio, org_name, inn_kpp, municipality) "
                "VALUES (%s, %s, %s, %s, %s) ON CONFLICT (user_id) DO UPDATE SET "
                "sender_fio = EXCLUDED.sender_fio, org_name = EXCLUDED.org_name, "
                "inn_kpp = EXCLUDED.inn_kpp, municipality = EXCLUDED.municipality;",
                (user_id, form_data['sender_fio'], form_data['org_name'], form_data['inn_kpp'], form_data['municipality'])
            )
        conn.commit()
        logger.info(f"Шаблон заявки для пользователя {user_id} сохранен.")
    except Exception as e:
        logger.error(f"Ошибка при сохранении шаблона заявки для {user_id}: {e}")
    finally:
        if conn: conn.close()

def load_akc_defaults(user_id: int) -> Optional[Dict[str, str]]:
    """Загружает шаблон данных для заявки АЦК."""
    conn = get_db_connection()
    if not conn: return None
    try:
        with conn.cursor() as cursor:
            cursor.execute("SELECT sender_fio, org_name, inn_kpp, municipality FROM akc_sender_defaults WHERE user_id = %s", (user_id,))
            result = cursor.fetchone()
        if result:
            return {
                'sender_fio': result[0],
                'org_name': result[1],
                'inn_kpp': result[2],
                'municipality': result[3]
            }
        return None
    finally:
        if conn: conn.close()

def save_akc_logins(user_id: int, logins: str):
    """Сохраняет логины пользователя для заявки АЦК."""
    conn = get_db_connection()
    if not conn: return
    try:
        with conn.cursor() as cursor:
            cursor.execute(
                "INSERT INTO akc_login_defaults (user_id, logins) VALUES (%s, %s) "
                "ON CONFLICT (user_id) DO UPDATE SET logins = EXCLUDED.logins;",
                (user_id, logins)
            )
        conn.commit()
        logger.info(f"Логины для пользователя {user_id} сохранены.")
    except Exception as e:
        logger.error(f"Ошибка при сохранении логинов для {user_id}: {e}")
    finally:
        if conn: conn.close()

def load_akc_logins(user_id: int) -> Optional[str]:
    """Загружает сохраненные логины пользователя."""
    conn = get_db_connection()
    if not conn: return None
    try:
        with conn.cursor() as cursor:
            cursor.execute("SELECT logins FROM akc_login_defaults WHERE user_id = %s", (user_id,))
            result = cursor.fetchone()
        return result[0] if result else None
    finally:
        if conn: conn.close()


# --- 4. ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ ---
def create_excel_report(cert_data_list: List[Dict[str, Any]], user_threshold: int) -> io.BytesIO:
    """Создает Excel-отчет на основе данных сертификатов."""
    wb = Workbook()
    ws = wb.active
    ws.title = "Отчет по сертификатам"
    ws.append(list(EXCEL_HEADERS))
    
    sorted_cert_data = sorted(cert_data_list, key=lambda x: x["Действителен до"])
    
    for cert_data in sorted_cert_data:
        row = [
            cert_data["ФИО"], 
            cert_data["Учреждение"], 
            cert_data["Серийный номер"], 
            cert_data["Действителен с"].strftime("%d.%m.%Y"), 
            cert_data["Действителен до"].strftime("%d.%m.%Y"), 
            cert_data["Осталось дней"]
        ]
        ws.append(row)
        
        last_row = ws.max_row
        days_left = cert_data["Осталось дней"]
        fill_color = None
        
        if days_left < 0:
            fill_color = RED_FILL
        elif 0 <= days_left <= user_threshold:
            fill_color = ORANGE_FILL
        else:
            fill_color = GREEN_FILL
            
        if fill_color:
            for cell in ws[last_row]:
                cell.fill = fill_color
                
    for column in ws.columns:
        max_length = 0
        column_letter = get_column_letter(column[0].column)
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = (max_length + 2)
        ws.column_dimensions[column_letter].width = adjusted_width
        
    excel_buffer = io.BytesIO()
    wb.save(excel_buffer)
    excel_buffer.seek(0)
    return excel_buffer

def generate_summary_message(cert_data_list: List[Dict[str, Any]], user_threshold: int) -> str:
    """Генерирует сводное сообщение о скоро истекающих сертификатах."""
    expiring_soon_certs = []
    for cert_data in cert_data_list:
        days_left = cert_data["Осталось дней"]
        if 0 <= days_left <= user_threshold:
            expiring_soon_certs.append(f"👤 {cert_data['ФИО']} — {cert_data['Действителен до'].strftime('%d.%m.%Y')} (осталось {days_left} дн.)")
            
    if expiring_soon_certs:
        message_parts = [f"⚠️ Скоро истекают ({user_threshold} дней):", *expiring_soon_certs]
        return "\n".join(message_parts)
    else:
        return "✅ Сертификатов, истекающих в ближайшее время, не найдено."

def get_certificate_info(cert_bytes: bytes) -> Optional[Dict[str, Any]]:
    """Извлекает информацию из файла сертификата."""
    try:
        try:
            cert = x509.load_pem_x509_certificate(cert_bytes, default_backend())
        except ValueError:
            cert = x509.load_der_x509_certificate(cert_bytes, default_backend())
            
        try:
            subject_common_name = cert.subject.get_attributes_for_oid(x509.OID_COMMON_NAME)[0].value
        except IndexError:
            subject_common_name = "Неизвестно"
            
        try:
            organization_name = cert.subject.get_attributes_for_oid(x509.OID_ORGANIZATION_NAME)[0].value
        except IndexError:
            organization_name = "Неизвестно"
            
        serial_number = f"{cert.serial_number:X}"
        valid_from = cert.not_valid_before.date()
        valid_until = cert.not_valid_after.date()
        days_left = (valid_until - datetime.now().date()).days
        
        return {
            "ФИО": subject_common_name,
            "Учреждение": organization_name,
            "Серийный номер": serial_number,
            "Действителен с": valid_from,
            "Действителен до": valid_until,
            "Осталось дней": days_left
        }
    except Exception as e:
        logger.error(f"Ошибка при парсинге сертификата: {e}")
        return None

def _process_file_content(file_bytes: bytes, file_name: str) -> List[Dict[str, Any]]:
    """Обрабатывает содержимое файла (сертификат или ZIP-архив)."""
    all_certs_data = []
    if file_name.lower().endswith(".zip"):
        try:
            with zipfile.ZipFile(io.BytesIO(file_bytes), 'r') as z:
                for member in z.namelist():
                    if member.lower().endswith(ALLOWED_EXTENSIONS):
                        with z.open(member) as cert_file:
                            cert_info = get_certificate_info(cert_file.read())
                            if cert_info:
                                all_certs_data.append(cert_info)
        except zipfile.BadZipFile:
            logger.error(f"Получен поврежденный ZIP-файл: {file_name}", exc_info=True)
            return []
    elif file_name.lower().endswith(ALLOWED_EXTENSIONS):
        cert_info = get_certificate_info(file_bytes)
        if cert_info:
            all_certs_data.append(cert_info)
    return all_certs_data

def create_akc_docx(form_data: dict) -> io.BytesIO:
    """Создает DOCX-файл заявки АЦК с несколькими записями."""
    doc = docx.Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Приложение 5 к Регламенту взаимодействия\nминистерства финансов Амурской области и\nУчастников юридически значимого\nэлектронного документооборота")
    font = run.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    doc.add_paragraph() 

    p = doc.add_paragraph()
    run = p.add_run("От кого: ")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run = p.add_run(f"{form_data.get('sender_fio', '')}\n")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    run = p.add_run("(Ф.И.О. представителя учреждения)\n")
    run.italic = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    
    run = p.add_run(f"{form_data.get('org_name', '')}\n")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    run = p.add_run("(наименование учреждения)\n")
    run.italic = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)

    run = p.add_run(f"{form_data.get('inn_kpp', '')}\n")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    run = p.add_run("(ИНН/КПП)\n")
    run.italic = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)

    run = p.add_run(f"{form_data.get('municipality', '')}\n")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    run = p.add_run("(наименование муниципального образования)\n")
    run.italic = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)

    run = p.add_run(f"{datetime.now().strftime('%d.%m.%Y')}\n")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    run = p.add_run("(дата)")
    run.italic = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ЗАЯВКА\nна регистрацию пользователя ЦИТП")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    
    headers = [
        "Субъект ЭП", "Роль субъекта в ЦИТП (Руководитель, Бухгалтер, Специалист ГИС ГМП)", 
        "Наименование ЦИТП (АЦК-Финансы, АЦК-Планирование)", 
        "Серийный номер сертификата", "Имя файла сертификата", 
        "Имя пользователя для входа в ЦИТП, под которым производится подписание", 
        "Действие(добавить, удалить, заменить, заблокировать)"
    ]
    
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = headers[i]
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)

    # Добавляем строки для каждого сертификата
    for cert_data in form_data.get('certificates', []):
        row_cells = table.add_row().cells
        row_cells[0].text = cert_data.get('cert_owner', '')
        row_cells[1].text = cert_data.get('role', '')
        row_cells[2].text = cert_data.get('citp_name', '')
        row_cells[3].text = cert_data.get('cert_serial', '')
        row_cells[4].text = cert_data.get('cert_filename', '')
        row_cells[5].text = cert_data.get('logins', '')
        row_cells[6].text = cert_data.get('action', '')
        for cell in row_cells:
            cell.paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()

    footer_table = doc.add_table(rows=4, cols=3)
    
    footer_table.cell(0, 0).text = "№ Записи в электронном журнале:"
    footer_table.cell(0, 1).text = "Статус:  выполнено/отказано"
    
    footer_table.cell(1, 0).text = "Заведение пользователя:"
    footer_table.cell(1, 1).text = "Подпись представителя учреждения"
    footer_table.cell(1, 2).text = "Дата"
    
    footer_table.cell(2, 0).text = "Дата:"
    footer_table.cell(2, 1).text = "Исполнитель:"
    
    footer_table.cell(3, 0).text = "Установка СКП ЭП"
    footer_table.cell(3, 1).text = "М.П."
    footer_table.cell(3, 2).text = "Номер телефона или e-mail"

    for row in footer_table.rows:
        for cell in row.cells:
            if cell.text:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)

    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer


# --- 5. ОБРАБОТЧИКИ КОМАНД, КНОПОК И ДИАЛОГОВ ---
async def get_my_id(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Отправляет пользователю его Telegram ID."""
    user_id = update.effective_user.id
    await update.message.reply_text(f"Ваш User ID: `{user_id}`", parse_mode='Markdown')

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Обработчик команды /start. Отображает главное меню в зависимости от прав."""
    user = update.effective_user
    user_id = user.id
    
    keyboard = []
    row1, row2, row3 = [], [], []
    
    if has_permission(user_id, "cert_analysis", context):
        row1.append("📜 Анализ сертификатов")
    if has_permission(user_id, "akc_form", context):
        row1.append("📄 Заявка АЦК")
    if row1:
        keyboard.append(row1)

    if has_permission(user_id, "youtube", context):
        row2.append("🎬 Скачивание с YouTube")
    if row2:
        keyboard.append(row2)
        
    if has_permission(user_id, "admin", context):
        row3.append("🔑 Управление доступом")
    row3.append("❓ Помощь")
    keyboard.append(row3)

    reply_markup = ReplyKeyboardMarkup(keyboard, resize_keyboard=True)
    start_message = (
        f"Привет, {user.mention_html()}! 👋\n\n"
        "Я — ваш многофункциональный помощник. Выберите доступное действие на клавиатуре."
    )
    await update.message.reply_html(start_message, reply_markup=reply_markup)
    return ConversationHandler.END

async def help_command(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Обработчик кнопки 'Помощь'. Отправляет справочную информацию."""
    help_text = (
        "Я могу помочь вам с несколькими задачами:\n\n"
        "📜 **Анализ сертификатов**\n"
        "Нажмите кнопку и отправьте файлы `.cer`, `.crt` или `.zip`-архив для создания Excel-отчета.\n\n"
        "📄 **Заявка АЦК**\n"
        "Нажмите кнопку, чтобы запустить пошаговый мастер создания заявки в формате DOCX.\n\n"
        "🎬 **Скачивание с YouTube**\n"
        "Нажмите кнопку и отправьте ссылку, чтобы скачать видео.\n\n"
        "🔑 **Управление доступом** (только для администраторов)\n"
        "Позволяет добавлять и удалять пользователей, а также настраивать их права."
    )
    await update.message.reply_text(help_text)

def download_video_sync(url: str, ydl_opts: dict) -> str:
    """Синхронная функция для скачивания видео с помощью yt-dlp."""
    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
        info = ydl.extract_info(url, download=True)
        return ydl.prepare_filename(info)

async def handle_youtube_link(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Обрабатывает полученную ссылку на YouTube видео."""
    url = update.message.text
    msg = await update.message.reply_text("Получаю информацию о видео...")
    
    try:
        ydl_opts_info = {'quiet': True, 'format': 'bestvideo[ext=mp4]+bestaudio[ext=m4a]/best[ext=mp4]/best'}
        with yt_dlp.YoutubeDL(ydl_opts_info) as ydl:
            info_dict = ydl.extract_info(url, download=False)
        
        filesize = info_dict.get('filesize') or info_dict.get('filesize_approx')
        title = info_dict.get('title', 'Без названия')
        
        if not filesize:
            await msg.edit_text("❌ Не удалось определить размер видео."); return ConversationHandler.END

        if filesize > MAX_VIDEO_SIZE_BYTES:
            size_in_mb = filesize / 1024 / 1024
            await msg.edit_text(f"❌ Видео '{title}' слишком большое ({size_in_mb:.1f} МБ)."); return ConversationHandler.END

        context.user_data['youtube_url'] = url
        context.user_data['youtube_title'] = title
        
        size_in_mb = filesize / 1024 / 1024
        keyboard = [[InlineKeyboardButton("✅ Да, скачать", callback_data='yt_confirm'), InlineKeyboardButton("❌ Нет, отмена", callback_data='yt_cancel')]]
        reply_markup = InlineKeyboardMarkup(keyboard)
        await msg.edit_text(f"**Название:** {title}\n**Размер:** {size_in_mb:.1f} МБ\n\nНачать скачивание?", reply_markup=reply_markup, parse_mode='Markdown')
        return CONFIRMING_DOWNLOAD

    except Exception as e:
        logger.error(f"Ошибка при получении информации о YouTube видео: {e}", exc_info=True)
        await msg.edit_text(f"❌ Не удалось получить информацию по ссылке: {url}"); return ConversationHandler.END

async def start_download_confirmed(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Начинает скачивание видео после подтверждения."""
    query = update.callback_query
    await query.answer()
    url = context.user_data.get('youtube_url')
    title = context.user_data.get('youtube_title', 'видео')
    user_id = update.effective_user.id

    if not url:
        await query.edit_message_text("❌ Произошла ошибка, начните заново."); return ConversationHandler.END

    await query.edit_message_text(f"Начинаю загрузку '{title}'...")
    
    ydl_opts = {'format': 'bestvideo[ext=mp4]+bestaudio[ext=m4a]/best[ext=mp4]/best', 'outtmpl': f'{uuid.uuid4()}.%(ext)s', 'quiet': True}
    
    try:
        video_filename = await asyncio.to_thread(download_video_sync, url, ydl_opts)
        await query.edit_message_text("Видео скачано. Отправляю...")
        with open(video_filename, 'rb') as video_file:
            await context.bot.send_video(chat_id=user_id, video=video_file, supports_streaming=True, read_timeout=120, write_timeout=120)
        os.remove(video_filename)
        await query.message.delete()
    except Exception as e:
        logger.error(f"Ошибка при скачивании/отправке видео: {e}", exc_info=True)
        await query.edit_message_text(f"❌ Не удалось обработать видео: {url}")
    
    return ConversationHandler.END

async def cancel_download(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Отменяет скачивание видео."""
    query = update.callback_query
    await query.answer()
    await query.edit_message_text("Скачивание отменено.")
    return ConversationHandler.END

async def youtube_entry(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Точка входа в диалог скачивания с YouTube."""
    if not has_permission(update.effective_user.id, "youtube", context): return ConversationHandler.END
    await update.message.reply_text("Пожалуйста, отправьте ссылку на YouTube видео.")
    return AWAITING_YOUTUBE_LINK

async def invalid_youtube_link(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Обрабатывает неверную ссылку на YouTube."""
    await update.message.reply_text("Это не похоже на ссылку YouTube. Пожалуйста, отправьте правильную ссылку или отмените действие.")
    return AWAITING_YOUTUBE_LINK

async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Отменяет текущий диалог."""
    await update.message.reply_text('Действие отменено.', reply_markup=ReplyKeyboardRemove())
    # Очищаем данные формы, чтобы избежать проблем при следующем запуске
    context.user_data.clear()
    return ConversationHandler.END

# --- ЛОГИКА ДИАЛОГА АНАЛИЗА СЕРТИФИКАТОВ ---
async def cert_analysis_start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Начинает диалог анализа сертификатов."""
    if not has_permission(update.effective_user.id, "cert_analysis", context): return ConversationHandler.END
    context.user_data['cert_analysis_data'] = {'files': []}
    keyboard = ReplyKeyboardMarkup([["Готово"]], resize_keyboard=True, one_time_keyboard=True)
    await update.message.reply_text(
        "**Анализ сертификатов** 📊\n\n"
        "Отправьте мне один или несколько файлов сертификатов (.cer, .crt, .zip).\n"
        "Когда закончите, нажмите кнопку **'Готово'**.",
        reply_markup=keyboard,
        parse_mode='Markdown'
    )
    return CERT_AWAIT_FILES

async def handle_cert_upload(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Собирает файлы для анализа."""
    document = update.message.document
    file_name = document.file_name
    file_id = document.file_id

    file_object = await context.bot.get_file(file_id)
    file_buffer = io.BytesIO()
    await file_object.download_to_memory(file_buffer)
    
    context.user_data['cert_analysis_data']['files'].append({
        'name': file_name,
        'bytes': file_buffer.getvalue()
    })
    
    await update.message.reply_text(f"✅ Файл `{file_name}` добавлен. Отправьте следующий или нажмите 'Готово'.", parse_mode='Markdown')
    return CERT_AWAIT_FILES

async def ask_for_threshold_confirmation(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Запрашивает подтверждение порога дней."""
    if not context.user_data['cert_analysis_data'].get('files'):
        await update.message.reply_text("Вы не добавили ни одного файла. Действие отменено.", reply_markup=ReplyKeyboardRemove())
        context.user_data.pop('cert_analysis_data', None)
        return ConversationHandler.END

    user_id = update.effective_user.id
    current_threshold = await get_user_threshold(user_id, context)
    
    keyboard = [
        [InlineKeyboardButton(f"✅ Использовать текущий ({current_threshold} дн.)", callback_data='cert_use_current')],
        [InlineKeyboardButton("✏️ Ввести новый", callback_data='cert_enter_new')]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await update.message.reply_text(
        "Файлы приняты. Теперь выберите порог дней для оповещения об истекающих сертификатах:",
        reply_markup=reply_markup
    )
    return CERT_AWAIT_THRESHOLD

async def prompt_for_new_threshold(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Запрашивает новый порог дней."""
    query = update.callback_query
    await query.answer()
    await query.edit_message_text("Введите новое число дней (например, 60).")
    return CERT_TYPING_THRESHOLD

async def process_with_threshold(update: Update, context: ContextTypes.DEFAULT_TYPE, threshold: int) -> int:
    """Общая функция для обработки и создания отчета."""
    message = update.message or update.callback_query.message
    await message.reply_text("Анализирую...", reply_markup=ReplyKeyboardRemove())
    
    all_certs_data = []
    for file_info in context.user_data['cert_analysis_data']['files']:
        processed_data = _process_file_content(file_info['bytes'], file_info['name'])
        all_certs_data.extend(processed_data)
        
    if not all_certs_data:
        await message.reply_text("Не удалось найти/проанализировать сертификаты в отправленных файлах.")
    else:
        excel_buffer = create_excel_report(all_certs_data, threshold)
        summary_message = generate_summary_message(all_certs_data, threshold)
        await message.reply_text(summary_message)
        await message.reply_document(document=excel_buffer, filename="Сертификаты_отчет.xlsx")
        
    context.user_data.pop('cert_analysis_data', None)
    return ConversationHandler.END

async def process_with_current_threshold(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Обрабатывает отчет с текущим порогом."""
    query = update.callback_query
    await query.answer()
    await query.delete_message()
    user_id = update.effective_user.id
    current_threshold = await get_user_threshold(user_id, context)
    return await process_with_threshold(update, context, current_threshold)

async def set_new_threshold_and_process(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Устанавливает новый порог и обрабатывает отчет."""
    user_id = update.effective_user.id
    try:
        new_threshold = int(update.message.text)
        if new_threshold <= 0:
            await update.message.reply_text("❌ Введите положительное число.")
            return CERT_TYPING_THRESHOLD
        
        save_user_threshold(user_id, new_threshold)
        context.user_data['threshold'] = new_threshold # Обновляем кэш
        await update.message.reply_html(f"✅ Порог изменен и сохранен: <b>{new_threshold}</b> дней.")
        return await process_with_threshold(update, context, new_threshold)

    except (ValueError):
        await update.message.reply_text("❌ Это не число. Отправьте, например: 60")
        return CERT_TYPING_THRESHOLD

# --- ЛОГИКА ДИАЛОГА ЗАЯВКИ АЦК ---
async def akc_start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Начинает диалог создания заявки АЦК и выводит описание."""
    if not has_permission(update.effective_user.id, "akc_form", context): return ConversationHandler.END
    user_id = update.effective_user.id
    context.user_data['akc_form'] = {
        'certificates': [] # Инициализируем список для сертификатов
    }

    description = (
        "**Мастер создания заявки АЦК** 📄\n\n"
        "Этот мастер поможет вам пошагово сформировать заявку на регистрацию пользователя ЦИТП.\n\n"
        "**Что он делает:**\n"
        "1. Запрашивает данные для шапки документа (ФИО, организация и т.д.).\n"
        "2. Сохраняет эти данные как шаблон для ускорения работы в будущем.\n"
        "3. Просит прикрепить **один или несколько** файлов сертификатов.\n"
        "4. Помогает поочередно настроить каждую запись в заявке.\n"
        "5. В итоге создает **ZIP-архив**, содержащий готовую заявку в формате **DOCX** и все прикрепленные сертификаты.\n\n"
        "------------------------------------\n\n"
    )

    defaults = load_akc_defaults(user_id)
    if defaults:
        context.user_data['akc_form'].update(defaults)
        text = description + (
            "Найдены сохраненные данные для шапки заявки:\n\n"
            f"• **От кого:** {defaults['sender_fio']}\n"
            f"• **Учреждение:** {defaults['org_name']}\n"
            f"• **ИНН/КПП:** {defaults['inn_kpp']}\n"
            f"• **МО:** {defaults['municipality']}\n\n"
            "Использовать эти данные?"
        )
        keyboard = [[InlineKeyboardButton("✅ Да, использовать", callback_data='akc_use_defaults')], [InlineKeyboardButton("✏️ Заполнить заново", callback_data='akc_refill')]]
        reply_markup = InlineKeyboardMarkup(keyboard)
        await update.message.reply_text(text, reply_markup=reply_markup, parse_mode='Markdown')
        return AKC_CONFIRM_DEFAULTS
    else:
        text = description + "Начинаем! Введите **ФИО представителя учреждения**:"
        await update.message.reply_text(text, parse_mode='Markdown')
        return AKC_SENDER_FIO

async def akc_use_defaults(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Использует сохраненный шаблон для шапки и запрашивает файлы."""
    query = update.callback_query
    await query.answer()
    await query.edit_message_text("Данные шапки применены.")
    
    keyboard = ReplyKeyboardMarkup([["Готово"]], resize_keyboard=True, one_time_keyboard=True)
    await query.message.reply_text(
        "Теперь отправьте мне **один или несколько** файлов сертификатов (.cer, .crt).\n"
        "Когда закончите, нажмите кнопку **'Готово'**.",
        reply_markup=keyboard
    )
    return AKC_AWAIT_CERTIFICATES

async def akc_refill_defaults(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Запускает процесс повторного заполнения шапки заявки."""
    query = update.callback_query
    await query.answer()
    await query.edit_message_text("Хорошо, давайте заполним данные заново.\n\nВведите **ФИО представителя учреждения**:", parse_mode='Markdown')
    return AKC_SENDER_FIO

async def akc_get_sender_fio(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    context.user_data['akc_form']['sender_fio'] = update.message.text
    await update.message.reply_text("Введите **полное наименование учреждения**:", parse_mode='Markdown')
    return AKC_ORG_NAME

async def akc_get_org_name(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    context.user_data['akc_form']['org_name'] = update.message.text
    await update.message.reply_text("Введите **ИНН/КПП** учреждения:", parse_mode='Markdown')
    return AKC_INN_KPP

async def akc_get_inn_kpp(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    context.user_data['akc_form']['inn_kpp'] = update.message.text
    await update.message.reply_text("Введите **наименование муниципального образования**:", parse_mode='Markdown')
    return AKC_MUNICIPALITY

async def akc_get_municipality(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Получает МО, сохраняет шаблон и запрашивает файлы."""
    user_id = update.effective_user.id
    context.user_data['akc_form']['municipality'] = update.message.text
    save_akc_defaults(user_id, context.user_data['akc_form'])
    await update.message.reply_text("Шапка заявки заполнена и сохранена.")
    
    keyboard = ReplyKeyboardMarkup([["Готово"]], resize_keyboard=True, one_time_keyboard=True)
    await update.message.reply_text(
        "Теперь отправьте мне **один или несколько** файлов сертификатов (.cer, .crt).\n"
        "Когда закончите, нажмите кнопку **'Готово'**.",
        reply_markup=keyboard
    )
    return AKC_AWAIT_CERTIFICATES

async def akc_add_certificate(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Обрабатывает и добавляет в список полученный файл сертификата."""
    document = update.message.document
    file_id = document.file_id
    
    try:
        file_object = await context.bot.get_file(file_id)
        file_buffer = io.BytesIO()
        await file_object.download_to_memory(file_buffer)
        cert_bytes = file_buffer.getvalue()
        
        cert_data = get_certificate_info(cert_bytes)
        
        if not cert_data:
            await update.message.reply_text(f"❌ Не удалось прочитать данные из файла `{document.file_name}`. Попробуйте другой файл.", parse_mode='Markdown')
            return AKC_AWAIT_CERTIFICATES

        context.user_data['akc_form']['certificates'].append({
            'cert_owner': cert_data['ФИО'],
            'cert_serial': cert_data['Серийный номер'],
            'cert_filename': document.file_name,
            'cert_bytes': cert_bytes,
            'role': '', 'citp_name': '', 'logins': '', 'action': ''
        })
        
        await update.message.reply_text(f"✅ Сертификат `{document.file_name}` добавлен. Отправьте следующий или нажмите 'Готово'.", parse_mode='Markdown')
        
    except Exception as e:
        logger.error(f"Ошибка при обработке файла сертификата для заявки: {e}", exc_info=True)
        await update.message.reply_text("❌ Произошла ошибка при обработке файла.")
        
    return AKC_AWAIT_CERTIFICATES

async def akc_start_data_loop(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Начинает цикл настройки для каждого добавленного сертификата."""
    if not context.user_data['akc_form'].get('certificates'):
        await update.message.reply_text("Вы не добавили ни одного сертификата. Действие отменено.", reply_markup=ReplyKeyboardRemove())
        context.user_data.pop('akc_form', None)
        return ConversationHandler.END

    await update.message.reply_text("Отлично! Все файлы приняты. Начинаем настройку каждой записи.", reply_markup=ReplyKeyboardRemove())
    context.user_data['akc_form']['cert_index'] = 0
    await _akc_ask_for_role(update, context)
    return AKC_ROLE

async def _akc_ask_for_role(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Вспомогательная функция для запроса роли для текущего сертификата."""
    cert_index = context.user_data['akc_form']['cert_index']
    cert_list = context.user_data['akc_form']['certificates']
    cert_data = cert_list[cert_index]
    cert_owner = cert_data['cert_owner']
    
    text = (
        f"➡️ **Настройка записи {cert_index + 1} из {len(cert_list)}**\n"
        f"Владелец: **{cert_owner}**\n\n"
        "Выберите **роль субъекта**:"
    )
    keyboard = [
        [InlineKeyboardButton("Руководитель", callback_data='role_Руководитель')],
        [InlineKeyboardButton("Бухгалтер", callback_data='role_Бухгалтер')],
        [InlineKeyboardButton("Специалист ГИС ГМП", callback_data='role_Специалист ГИС ГМП')]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)

    if update.callback_query:
        await update.callback_query.edit_message_text(text, reply_markup=reply_markup, parse_mode='Markdown')
    else:
        await update.message.reply_text(text, reply_markup=reply_markup, parse_mode='Markdown')

async def akc_get_role(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Получает роль для текущего сертификата."""
    query = update.callback_query
    await query.answer()
    role = query.data.split('_')[1]
    
    cert_index = context.user_data['akc_form']['cert_index']
    context.user_data['akc_form']['certificates'][cert_index]['role'] = role

    keyboard = [[InlineKeyboardButton("АЦК-Финансы", callback_data='citp_АЦК-Финансы')], [InlineKeyboardButton("АЦК-Планирование", callback_data='citp_АЦК-Планирование')]]
    reply_markup = InlineKeyboardMarkup(keyboard)
    await query.edit_message_text(text=f"Выбрана роль: {role}.\n\nВыберите **Наименование ЦИТП**:", reply_markup=reply_markup, parse_mode='Markdown')
    return AKC_CITP_NAME

async def akc_get_citp_name(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Получает ЦИТП и проверяет сохраненные логины."""
    query = update.callback_query
    await query.answer()
    citp_name = query.data.split('_')[1]

    cert_index = context.user_data['akc_form']['cert_index']
    context.user_data['akc_form']['certificates'][cert_index]['citp_name'] = citp_name
    
    user_id = update.effective_user.id
    saved_logins = load_akc_logins(user_id)
    
    if saved_logins:
        context.user_data['akc_saved_logins'] = saved_logins
        text = (
            f"Выбрана система: {citp_name}.\n\n"
            f"Найдены ранее сохраненные логины: `{saved_logins}`\n\n"
            "Использовать их?"
        )
        keyboard = [
            [InlineKeyboardButton("✅ Да, использовать", callback_data='logins_use_saved')],
            [InlineKeyboardButton("✏️ Ввести новые", callback_data='logins_enter_new')]
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        await query.edit_message_text(text=text, reply_markup=reply_markup, parse_mode='Markdown')
        return AKC_CONFIRM_LOGINS
    else:
        await query.edit_message_text(text=f"Выбрана система: {citp_name}.\n\nВведите **имена пользователей (логины)**, через запятую:", parse_mode='Markdown')
        return AKC_LOGINS

async def akc_use_saved_logins(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Использует сохраненные логины."""
    query = update.callback_query
    await query.answer()
    
    saved_logins = context.user_data.get('akc_saved_logins')
    cert_index = context.user_data['akc_form']['cert_index']
    context.user_data['akc_form']['certificates'][cert_index]['logins'] = saved_logins
    
    keyboard = [[InlineKeyboardButton(text, callback_data=f'action_{text}') for text in ["Добавить", "Удалить"]], [InlineKeyboardButton(text, callback_data=f'action_{text}') for text in ["Заменить", "Заблокировать"]]]
    reply_markup = InlineKeyboardMarkup(keyboard)
    await query.edit_message_text("Логины применены.\n\nВыберите **действие** с сертификатом:", reply_markup=reply_markup, parse_mode='Markdown')
    
    context.user_data.pop('akc_saved_logins', None)
    return AKC_ACTION

async def akc_enter_new_logins(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Запрашивает ввод новых логинов."""
    query = update.callback_query
    await query.answer()
    await query.edit_message_text("Введите **новые имена пользователей (логины)**, через запятую:", parse_mode='Markdown')
    return AKC_LOGINS

async def akc_get_logins(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Получает и сохраняет новые логины."""
    user_id = update.effective_user.id
    logins = update.message.text
    
    cert_index = context.user_data['akc_form']['cert_index']
    context.user_data['akc_form']['certificates'][cert_index]['logins'] = logins
    save_akc_logins(user_id, logins)
    
    keyboard = [[InlineKeyboardButton(text, callback_data=f'action_{text}') for text in ["Добавить", "Удалить"]], [InlineKeyboardButton(text, callback_data=f'action_{text}') for text in ["Заменить", "Заблокировать"]]]
    reply_markup = InlineKeyboardMarkup(keyboard)
    await update.message.reply_text("Новые логины сохранены.\n\nВыберите **действие** с сертификатом:", reply_markup=reply_markup, parse_mode='Markdown')
    return AKC_ACTION

async def akc_get_action(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Получает действие и решает, продолжать цикл или завершать."""
    query = update.callback_query
    await query.answer()
    action = query.data.split('_')[1]

    cert_index = context.user_data['akc_form']['cert_index']
    context.user_data['akc_form']['certificates'][cert_index]['action'] = action

    cert_index += 1
    context.user_data['akc_form']['cert_index'] = cert_index

    if cert_index < len(context.user_data['akc_form']['certificates']):
        await _akc_ask_for_role(update, context)
        return AKC_ROLE
    else:
        await query.edit_message_text("Все записи настроены. Формирую итоговый ZIP-архив...")
        await akc_finish(update, context)
        context.user_data.pop('akc_form', None)
        return ConversationHandler.END

async def akc_finish(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Завершает создание заявки, формирует и отправляет ZIP-архив."""
    form_data = context.user_data['akc_form']
    try:
        docx_buffer = create_akc_docx(form_data)
        docx_filename = f"Заявка_АЦК_{form_data.get('sender_fio', 'пользователь')}.docx"
        
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w') as zip_file:
            zip_file.writestr(docx_filename, docx_buffer.getvalue())
            for cert_data in form_data.get('certificates', []):
                if cert_data.get('cert_bytes') and cert_data.get('cert_filename'):
                    zip_file.writestr(cert_data['cert_filename'], cert_data['cert_bytes'])
        zip_buffer.seek(0)
        
        zip_filename = f"Заявка_АЦК_{form_data.get('sender_fio', 'пользователь')}.zip"
        
        await context.bot.send_document(chat_id=update.effective_chat.id, document=zip_buffer, filename=zip_filename, caption="✅ Ваша заявка и сертификат в ZIP-архиве.")
        
    except Exception as e:
        logger.error(f"Ошибка при создании или отправке ZIP-архива: {e}", exc_info=True)
        await context.bot.send_message(chat_id=update.effective_chat.id, text="❌ Произошла ошибка при создании архива.")

# --- ЛОГИКА ДИАЛОГА УПРАВЛЕНИЯ ДОСТУПОМ ---
async def access_management_start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Начинает диалог управления доступом."""
    if not has_permission(update.effective_user.id, "admin", context): return ConversationHandler.END
    await _show_access_menu(update, context)
    return ACCESS_MENU

async def _show_access_menu(update: Update, context: ContextTypes.DEFAULT_TYPE, message_id: int = None):
    """Отображает меню управления доступом со списком пользователей."""
    permissions_dict = context.bot_data.get('permissions', {})
    
    text_lines = ["**🔑 Управление доступом**\n\nТекущие пользователи и их права:"]
    user_list_empty = True
    for user_id, user_data in permissions_dict.items():
        user_list_empty = False
        perms = user_data.get('perms', set())
        name = user_data.get('name', f'Пользователь {user_id}')
        perms_str = ", ".join([AVAILABLE_PERMISSIONS.get(p, p) for p in perms])
        text_lines.append(f"• {name} (`{user_id}`): {perms_str}")

    if user_list_empty:
        text_lines.append("Нет пользователей с настроенными правами.")

    keyboard = [
        [InlineKeyboardButton("➕ Добавить пользователя", callback_data='access_add')],
        [InlineKeyboardButton("❌ Удалить пользователя", callback_data='access_delete')],
        [InlineKeyboardButton("⬅️ Назад в главное меню", callback_data='access_back')]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    message = update.message or update.callback_query.message
    if message_id:
        await context.bot.edit_message_text(chat_id=message.chat_id, message_id=message_id, text="\n".join(text_lines), reply_markup=reply_markup, parse_mode='Markdown')
    else:
        await message.reply_text("\n".join(text_lines), reply_markup=reply_markup, parse_mode='Markdown')

async def access_back(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Возвращает в главное меню."""
    query = update.callback_query
    await query.answer()
    await query.edit_message_text("Возврат в главное меню.")
    return ConversationHandler.END

async def prompt_add_user(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Запрашивает способ добавления нового пользователя."""
    query = update.callback_query
    await query.answer()
    await query.edit_message_text(
        "**Способ добавления пользователя:**\n\n"
        "1. **Переслать сообщение** - самый простой способ, бот автоматически определит ID и имя.\n"
        "2. **Ввести ID вручную** - если у вас есть только ID пользователя.\n\n"
        "Пожалуйста, перешлите сообщение или отправьте ID.",
        parse_mode='Markdown'
    )
    return AWAITING_USER_INFO

async def get_user_info(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Обрабатывает пересланное сообщение или введенный ID."""
    user_to_add = None
    user_id = None
    user_name = None

    # Проверяем, переслано ли сообщение
    if update.message.forward_origin and isinstance(update.message.forward_origin, MessageOriginUser):
        user_to_add = update.message.forward_origin.sender_user
        user_id = user_to_add.id
        user_name = user_to_add.full_name
    # Если нет, пытаемся обработать как ID
    else:
        try:
            user_id = int(update.message.text)
            user_name = f"Пользователь {user_id}" # Имя по умолчанию для ручного ввода
        except (ValueError, TypeError):
            await update.message.reply_text("Не удалось распознать. Пожалуйста, либо перешлите сообщение, либо введите корректный ID.")
            return AWAITING_USER_INFO
            
    context.user_data['new_user_id'] = user_id
    context.user_data['new_user_name'] = user_name
    context.user_data['new_user_perms'] = set()
    
    await _show_permission_selection(update, context)
    return AWAITING_PERMISSIONS

async def _show_permission_selection(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Отображает клавиатуру для выбора разрешений."""
    new_user_id = context.user_data['new_user_id']
    new_user_name = context.user_data['new_user_name']
    selected_perms = context.user_data['new_user_perms']
    
    keyboard = []
    for perm_key, perm_name in AVAILABLE_PERMISSIONS.items():
        is_selected = "✅" if perm_key in selected_perms else "☑️"
        keyboard.append([InlineKeyboardButton(f"{is_selected} {perm_name}", callback_data=f"perm_{perm_key}")])
    
    keyboard.append([InlineKeyboardButton("💾 Сохранить", callback_data="perm_save")])
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    text = f"Выберите разрешения для **{new_user_name}** (`{new_user_id}`):"
    
    message = update.message or update.callback_query.message
    # Удаляем предыдущее сообщение (с просьбой переслать) и отправляем новое
    await message.delete()
    await message.chat.send_message(text, reply_markup=reply_markup, parse_mode='Markdown')

async def toggle_permission(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Переключает разрешение для нового пользователя."""
    query = update.callback_query
    await query.answer()
    
    perm_key = query.data.split('_', 1)[1]
    selected_perms = context.user_data['new_user_perms']
    
    if perm_key in selected_perms:
        selected_perms.remove(perm_key)
    else:
        selected_perms.add(perm_key)
        
    await _show_permission_selection(update, context)
    return AWAITING_PERMISSIONS

async def save_new_user(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Сохраняет нового пользователя и его права."""
    query = update.callback_query
    await query.answer()
    
    user_id = context.user_data['new_user_id']
    username = context.user_data['new_user_name']
    permissions = context.user_data['new_user_perms']
    
    if not permissions:
        await query.edit_message_text("Вы не выбрали ни одного разрешения. Добавление отменено.")
    else:
        db_save_user_permissions(user_id, username, permissions)
        context.bot_data['permissions'] = db_load_all_permissions() # Обновляем кэш
        await query.edit_message_text(f"Пользователь **{username}** (`{user_id}`) успешно добавлен/обновлен.", parse_mode='Markdown')

    context.user_data.clear()
    
    await _show_access_menu(update, context, message_id=query.message.message_id)
    return ACCESS_MENU

async def prompt_delete_user(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Показывает список пользователей для удаления."""
    query = update.callback_query
    await query.answer()
    
    permissions_dict = context.bot_data.get('permissions', {})
    keyboard = []
    
    for user_id, user_data in permissions_dict.items():
        if user_id == ADMIN_USER_ID: continue
        name = user_data.get('name', user_id)
        keyboard.append([InlineKeyboardButton(f"Удалить {name} (`{user_id}`)", callback_data=f"del_{user_id}")])
        
    if not keyboard:
        await query.edit_message_text("Нет пользователей для удаления.")
        await _show_access_menu(update, context, message_id=query.message.message_id)
        return ACCESS_MENU
        
    keyboard.append([InlineKeyboardButton("⬅️ Назад", callback_data='access_show_menu')])
    reply_markup = InlineKeyboardMarkup(keyboard)
    await query.edit_message_text("Выберите пользователя для удаления:", reply_markup=reply_markup, parse_mode='Markdown')
    return AWAITING_USER_TO_DELETE

async def delete_user(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Удаляет выбранного пользователя."""
    query = update.callback_query
    await query.answer()
    
    user_id_to_delete = int(query.data.split('_')[1])
    
    db_delete_user(user_id_to_delete)
    context.bot_data['permissions'] = db_load_all_permissions() # Обновляем кэш
    
    await query.edit_message_text(f"Пользователь `{user_id_to_delete}` удален.", parse_mode='Markdown')
    
    await _show_access_menu(update, context, message_id=query.message.message_id)
    return ACCESS_MENU

async def return_to_access_menu(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Возвращает к главному меню управления доступом."""
    query = update.callback_query
    await query.answer()
    await _show_access_menu(update, context, message_id=query.message.message_id)
    return ACCESS_MENU

# --- 6. ОСНОВНАЯ ФУНКЦИЯ ЗАПУСКА ---
async def main() -> None:
    """Главная функция для настройки и запуска бота."""
    if not TELEGRAM_BOT_TOKEN or not DATABASE_URL:
        logger.error("Не найден токен или URL базы данных. Проверьте переменные окружения.")
        return
        
    init_database()
    
    application = Application.builder().token(TELEGRAM_BOT_TOKEN).build()
    
    # Загружаем разрешения в кэш при старте
    application.bot_data['permissions'] = db_load_all_permissions()
    
    # Создаем фильтр на основе загруженных пользователей
    authorized_user_filter = filters.User(user_id=application.bot_data['permissions'].keys())

    cancel_handler = MessageHandler(filters.Regex('^/cancel$') | filters.Regex('^Отмена$'), cancel)
    
    cert_analysis_conv_handler = ConversationHandler(
        entry_points=[
            CommandHandler("cert", cert_analysis_start),
            MessageHandler(filters.Regex('^📜 Анализ сертификатов$'), cert_analysis_start)
        ],
        states={
            CERT_AWAIT_FILES: [
                MessageHandler(filters.Document.ALL, handle_cert_upload),
                MessageHandler(filters.Regex('^Готово$'), ask_for_threshold_confirmation)
            ],
            CERT_AWAIT_THRESHOLD: [
                CallbackQueryHandler(process_with_current_threshold, pattern='^cert_use_current$'),
                CallbackQueryHandler(prompt_for_new_threshold, pattern='^cert_enter_new$')
            ],
            CERT_TYPING_THRESHOLD: [MessageHandler(filters.TEXT & ~filters.COMMAND, set_new_threshold_and_process)],
        },
        fallbacks=[cancel_handler],
    )

    youtube_conv_handler = ConversationHandler(
        entry_points=[MessageHandler(filters.Regex('^🎬 Скачивание с YouTube$'), youtube_entry)],
        states={
            AWAITING_YOUTUBE_LINK: [MessageHandler(filters.Regex(YOUTUBE_URL_PATTERN), handle_youtube_link)],
            CONFIRMING_DOWNLOAD: [CallbackQueryHandler(start_download_confirmed, pattern='^yt_confirm$'), CallbackQueryHandler(cancel_download, pattern='^yt_cancel$')]
        },
        fallbacks=[cancel_handler]
    )
    
    akc_cert_filter = filters.Document.FileExtension("cer") | filters.Document.FileExtension("crt")
    
    akc_conv_handler = ConversationHandler(
        entry_points=[MessageHandler(filters.Regex('^📄 Заявка АЦК$'), akc_start)],
        states={
            AKC_CONFIRM_DEFAULTS: [CallbackQueryHandler(akc_use_defaults, pattern='^akc_use_defaults$'), CallbackQueryHandler(akc_refill_defaults, pattern='^akc_refill$')],
            AKC_SENDER_FIO: [MessageHandler(filters.TEXT & ~filters.COMMAND, akc_get_sender_fio)],
            AKC_ORG_NAME: [MessageHandler(filters.TEXT & ~filters.COMMAND, akc_get_org_name)],
            AKC_INN_KPP: [MessageHandler(filters.TEXT & ~filters.COMMAND, akc_get_inn_kpp)],
            AKC_MUNICIPALITY: [MessageHandler(filters.TEXT & ~filters.COMMAND, akc_get_municipality)],
            AKC_AWAIT_CERTIFICATES: [
                MessageHandler(akc_cert_filter, akc_add_certificate),
                MessageHandler(filters.Regex('^Готово$'), akc_start_data_loop)
            ],
            AKC_ROLE: [CallbackQueryHandler(akc_get_role, pattern='^role_')],
            AKC_CITP_NAME: [CallbackQueryHandler(akc_get_citp_name, pattern='^citp_')],
            AKC_CONFIRM_LOGINS: [
                CallbackQueryHandler(akc_use_saved_logins, pattern='^logins_use_saved$'),
                CallbackQueryHandler(akc_enter_new_logins, pattern='^logins_enter_new$')
            ],
            AKC_LOGINS: [MessageHandler(filters.TEXT & ~filters.COMMAND, akc_get_logins)],
            AKC_ACTION: [CallbackQueryHandler(akc_get_action, pattern='^action_')],
        },
        fallbacks=[cancel_handler],
        per_message=False
    )

    access_management_conv = ConversationHandler(
        entry_points=[MessageHandler(filters.Regex('^🔑 Управление доступом$'), access_management_start)],
        states={
            ACCESS_MENU: [
                CallbackQueryHandler(prompt_add_user, pattern='^access_add$'),
                CallbackQueryHandler(prompt_delete_user, pattern='^access_delete$'),
                CallbackQueryHandler(access_back, pattern='^access_back$'),
            ],
            AWAITING_USER_INFO: [
                MessageHandler(filters.FORWARDED, get_user_info),
                MessageHandler(filters.TEXT & ~filters.COMMAND, get_user_info)
            ],
            AWAITING_PERMISSIONS: [
                CallbackQueryHandler(save_new_user, pattern='^perm_save$'),
                CallbackQueryHandler(toggle_permission, pattern='^perm_'),
            ],
            AWAITING_USER_TO_DELETE: [
                CallbackQueryHandler(delete_user, pattern='^del_'),
                CallbackQueryHandler(return_to_access_menu, pattern='^access_show_menu$'),
            ]
        },
        fallbacks=[cancel_handler]
    )

    application.add_handler(cert_analysis_conv_handler)
    application.add_handler(youtube_conv_handler)
    application.add_handler(akc_conv_handler)
    application.add_handler(access_management_conv)
    
    application.add_handler(CommandHandler("my_id", get_my_id))
    application.add_handler(CommandHandler("start", start, filters=authorized_user_filter))
    
    application.add_handler(MessageHandler(filters.Regex("^(❓ Помощь)$") & authorized_user_filter, help_command))

    logger.info("Запускаю бота...")
    async with application:
        await application.start()
        await application.updater.start_polling(allowed_updates=Update.ALL_TYPES)
        await asyncio.Future()


# --- 7. ТОЧКА ВХОДА ---
if __name__ == "__main__":
    try:
        asyncio.run(main())
    except (KeyboardInterrupt, SystemExit):
        logger.info("Бот остановлен пользователем.")
